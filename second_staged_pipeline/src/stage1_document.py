"""Stage 1 — Document Processing.

Converts PDF/DOCX/TXT into structured pages and section-tagged chunks.
Uses PyMuPDF as primary PDF extractor (with pdfplumber fallback).
No OCR — documents are digital.
"""
from __future__ import annotations

import pathlib
import re
from dataclasses import dataclass, field

try:
    import fitz  # PyMuPDF
    _FITZ_AVAILABLE = True
except ImportError:
    _FITZ_AVAILABLE = False

try:
    import pdfplumber
    _PDFPLUMBER_AVAILABLE = True
except ImportError:
    _PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    from langdetect import detect as langdetect_detect
    _LANGDETECT_AVAILABLE = True
except ImportError:
    _LANGDETECT_AVAILABLE = False


_HEADING_RE = re.compile(
    r"^(#{1,4}\s+.+|[A-ZÀÂÇÉÈÊËÎÏÔÙÛÜÆŒ؀-ۿ]{3,}[\w\s؀-ۿ,.:()-]{0,80})$",
    re.MULTILINE,
)
_SECTION_SPLIT = re.compile(r"\n{2,}")
_MIN_CHUNK_CHARS = 80


def _log(msg: str):
    print(f"  [stage1] {msg}", flush=True)


@dataclass
class DocumentPage:
    page_num: int
    text: str
    tables: list[list[list[str]]] = field(default_factory=list)


@dataclass
class DocumentChunk:
    index: int
    text: str
    section_title: str | None = None
    page_num: int | None = None
    language: str | None = None


@dataclass
class ProcessedDocument:
    text: str
    pages: list[DocumentPage]
    chunks: list[DocumentChunk]
    language: str | None
    parser: str
    total_chars: int
    total_chunks: int


def process_document(path: str) -> ProcessedDocument:
    """Stage 1 entry point — parse, clean, detect language, chunk."""
    p = pathlib.Path(path)
    suffix = p.suffix.lower()

    if suffix == ".pdf":
        pages = _extract_pdf(path)
        parser = "pymupdf" if _FITZ_AVAILABLE else "pdfplumber"
    elif suffix in (".docx", ".doc"):
        pages = _extract_docx(path)
        parser = "python-docx"
    elif suffix == ".txt":
        pages = _extract_txt(path)
        parser = "txt"
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    full_text = _merge_pages(pages)
    full_text = _clean_text(full_text)
    language = _detect_language(full_text)
    chunks = _build_chunks(pages, full_text)

    _log(
        f"Parsed via {parser} — {len(pages)} pages, "
        f"{len(full_text)} chars, {len(chunks)} chunks, lang={language}"
    )

    return ProcessedDocument(
        text=full_text,
        pages=pages,
        chunks=chunks,
        language=language,
        parser=parser,
        total_chars=len(full_text),
        total_chunks=len(chunks),
    )


def _extract_pdf(path: str) -> list[DocumentPage]:
    if _FITZ_AVAILABLE:
        return _extract_pdf_fitz(path)
    if _PDFPLUMBER_AVAILABLE:
        return _extract_pdf_pdfplumber(path)
    raise ImportError("Neither PyMuPDF (fitz) nor pdfplumber is installed")


def _extract_pdf_fitz(path: str) -> list[DocumentPage]:
    _log("PDF extractor: PyMuPDF")
    pages: list[DocumentPage] = []
    doc = fitz.open(path)
    total = len(doc)
    for i, page in enumerate(doc, 1):
        print(f"\r  [stage1] PyMuPDF: page {i}/{total}", end="", flush=True)
        text = page.get_text("text")
        pages.append(DocumentPage(page_num=i, text=text or ""))
    print()
    doc.close()
    _log(f"PyMuPDF OK — {len(pages)} pages")
    return pages


def _extract_pdf_pdfplumber(path: str) -> list[DocumentPage]:
    _log("PDF extractor: pdfplumber (fallback)")
    pages: list[DocumentPage] = []
    with pdfplumber.open(path) as pdf:
        total = len(pdf.pages)
        for i, page in enumerate(pdf.pages, 1):
            print(f"\r  [stage1] pdfplumber: page {i}/{total}", end="", flush=True)
            text = page.extract_text() or ""
            raw_tables = page.extract_tables() or []
            tables = [t for t in raw_tables if t]
            pages.append(DocumentPage(page_num=i, text=text, tables=tables))
    print()
    _log(f"pdfplumber OK — {len(pages)} pages")
    return pages


def _extract_docx(path: str) -> list[DocumentPage]:
    if not _DOCX_AVAILABLE:
        raise ImportError("python-docx is not installed")
    _log("DOCX extractor: python-docx")
    doc = DocxDocument(path)

    headings: list[str] = []
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        if para.style.name.startswith("Heading"):
            headings.append(t)
            paragraphs.append(f"## {t}")
        else:
            paragraphs.append(t)

    text = "\n\n".join(paragraphs)
    _log(f"python-docx OK — {len(paragraphs)} paragraphs, {len(headings)} headings")
    return [DocumentPage(page_num=1, text=text)]


def _extract_txt(path: str) -> list[DocumentPage]:
    text = pathlib.Path(path).read_text(encoding="utf-8", errors="replace")
    return [DocumentPage(page_num=1, text=text)]


def _merge_pages(pages: list[DocumentPage]) -> str:
    return "\n\n".join(p.text for p in pages if p.text.strip())


def _clean_text(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    text = re.sub(r"(\s*\n\s*){3,}", "\n\n", text)
    return text.strip()


def _build_chunks(pages: list[DocumentPage], full_text: str) -> list[DocumentChunk]:
    """Split into paragraph-level chunks with section tracking."""
    paragraphs = _SECTION_SPLIT.split(full_text)
    chunks: list[DocumentChunk] = []
    current_section: str | None = None
    index = 0

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        if _is_heading(para):
            current_section = re.sub(r"^#{1,4}\s*", "", para).strip()
            continue

        if len(para) < _MIN_CHUNK_CHARS and chunks:
            prev = chunks[-1]
            chunks[-1] = DocumentChunk(
                index=prev.index,
                text=prev.text + "\n\n" + para,
                section_title=prev.section_title,
                page_num=prev.page_num,
                language=prev.language,
            )
            continue

        lang = _detect_language(para[:500]) if len(para) > 50 else None
        chunks.append(DocumentChunk(
            index=index,
            text=para,
            section_title=current_section,
            page_num=_guess_page(para, pages),
            language=lang,
        ))
        index += 1

    return chunks


def _is_heading(para: str) -> bool:
    stripped = para.strip()
    if re.match(r"^#{1,4}\s+", stripped):
        return True
    words = stripped.split()
    if len(words) <= 12 and stripped.upper() == stripped and len(stripped) < 100:
        return True
    return False


def _guess_page(text: str, pages: list[DocumentPage]) -> int | None:
    sample = text[:100]
    for page in pages:
        if sample in page.text:
            return page.page_num
    return None


def _detect_language(text: str) -> str | None:
    if not _LANGDETECT_AVAILABLE or not text.strip():
        return None
    try:
        return langdetect_detect(text[:2000])
    except Exception:
        return None
