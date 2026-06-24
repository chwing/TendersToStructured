import pathlib
from dataclasses import dataclass, field

try:
    import pdfplumber
    _PDFPLUMBER_AVAILABLE = True
except ImportError:
    _PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


def _log(msg: str):
    print(f"  [ingest] {msg}", flush=True)


@dataclass
class IngestedDocument:
    text: str
    pages: list[str] = field(default_factory=list)
    tables: list[str] = field(default_factory=list)
    parser: str = "unknown"


def ingest(path: str) -> IngestedDocument:
    suffix = path.lower()
    if suffix.endswith(".pdf"):
        return _ingest_pdf(path)
    elif suffix.endswith(".docx") or suffix.endswith(".doc"):
        return _ingest_docx(path)
    elif suffix.endswith(".txt"):
        return _ingest_txt(path)
    else:
        raise ValueError(f"Unsupported file type: {path}")


def _ingest_pdf(path: str) -> IngestedDocument:
    return _ingest_pdf_pdfplumber(path)


def _ingest_pdf_pdfplumber(path: str) -> IngestedDocument:
    if not _PDFPLUMBER_AVAILABLE:
        raise ImportError("pdfplumber is not installed")
    _log("Using pdfplumber ...")
    pages = []
    with pdfplumber.open(path) as pdf:
        total = len(pdf.pages)
        for i, page in enumerate(pdf.pages, 1):
            print(f"\r  [ingest] pdfplumber: page {i}/{total}", end="", flush=True)
            text = page.extract_text()
            if text:
                pages.append(text)
    print()  # newline after progress
    full_text = "\n\n".join(pages)
    _log(f"pdfplumber OK — {len(pages)} pages, {len(full_text)} chars")
    return IngestedDocument(text=full_text, pages=pages, tables=[], parser="pdfplumber")


def _ingest_txt(path: str) -> IngestedDocument:
    text = pathlib.Path(path).read_text(encoding="utf-8")
    _log(f"txt OK — {len(text)} chars")
    return IngestedDocument(text=text, pages=[text], tables=[], parser="txt")


def _ingest_docx(path: str) -> IngestedDocument:
    if not _DOCX_AVAILABLE:
        raise ImportError("python-docx is not installed")
    _log("Using python-docx ...")
    doc = DocxDocument(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    _log(f"python-docx OK — {len(paragraphs)} paragraphs, {len(text)} chars")
    return IngestedDocument(text=text, pages=[text], tables=[], parser="python-docx")
